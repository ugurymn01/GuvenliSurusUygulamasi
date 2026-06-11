# -*- coding: utf-8 -*-
"""SafeDrive proje raporu (.docx) üreteci — klasik akademik biçim."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---- Genel stil: Times New Roman 12, 1.5 satır, iki yana yaslı ----
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
pf = normal.paragraph_format
pf.line_spacing = 1.5
pf.space_after = Pt(6)

HEAD = RGBColor(0x1F, 0x38, 0x64)   # koyu lacivert (Word standart "Dark Blue")
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY = RGBColor(0x80, 0x80, 0x80)

sekil_sayac = [0]
sekil_listesi = []


def _set_font(run, name='Times New Roman'):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)


def _bottom_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), '1F3864')
    pBdr.append(bottom)
    pPr.append(pBdr)


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = HEAD
    _set_font(r)
    _bottom_border(p)
    return p


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12.5)
    r.font.color.rgb = BLACK
    _set_font(r)
    return p


def para(text, bold=False, italic=False, size=12, justify=True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    _set_font(r)
    return p


def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    _set_font(r)
    return p


def numbered(text):
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(text)
    _set_font(r)
    return p


def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
    r.font.size = Pt(10)
    return p


def _shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def tablo(basliklar, satirlar):
    t = doc.add_table(rows=1, cols=len(basliklar))
    t.style = 'Table Grid'           # sade siyah çerçeveli klasik tablo
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, b in enumerate(basliklar):
        hdr[i].text = ''
        _shade_cell(hdr[i], 'D9D9D9')
        run = hdr[i].paragraphs[0].add_run(b)
        run.bold = True
        run.font.size = Pt(10.5)
        _set_font(run)
    for satir in satirlar:
        cells = t.add_row().cells
        for i, deger in enumerate(satir):
            cells[i].text = ''
            run = cells[i].paragraphs[0].add_run(str(deger))
            run.font.size = Pt(10)
            _set_font(run)
    doc.add_paragraph()
    return t


def _set_row_height(row, cm):
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(cm * 567)))  # cm -> twips
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)


def sekil(aciklama):
    """Akademik 'Şekil' çerçeveli görsel kutusu + altında başlık."""
    sekil_sayac[0] += 1
    n = sekil_sayac[0]
    sekil_listesi.append((n, aciklama))
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_row_height(t.rows[0], 5.0)
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = cp.add_run('[ Görsel buraya eklenecek ]')
    rr.italic = True
    rr.font.size = Pt(10)
    rr.font.color.rgb = GREY
    _set_font(rr)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r1 = cap.add_run(f'Şekil {n}. ')
    r1.bold = True
    r1.font.size = Pt(10)
    _set_font(r1)
    r2 = cap.add_run(aciklama)
    r2.font.size = Pt(10)
    _set_font(r2)
    return t


# ============================ KAPAK ============================
for _ in range(5):
    doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('SafeDrive')
r.bold = True
r.font.size = Pt(36)
r.font.color.rgb = HEAD
_set_font(r)

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run('Güvenli Sürüş ve Sürücü Davranış Analizi Platformu')
r.font.size = Pt(16)
r.font.color.rgb = BLACK
_set_font(r)

s2 = doc.add_paragraph()
s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s2.add_run('Lojistik Filolar için Gerçek Zamanlı Telematik ve Filo Takip Sistemi')
r.italic = True
r.font.size = Pt(12)
r.font.color.rgb = GREY
_set_font(r)

for _ in range(8):
    doc.add_paragraph()
k = doc.add_paragraph()
k.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in ['PROJE RAPORU', '[DÜZENLE: Ders / Bölüm / Üniversite adı]',
             '[DÜZENLE: Hazırlayan(lar)ın adı-soyadı, numarası]', 'Haziran 2026']:
    rr = k.add_run(line + '\n')
    rr.font.size = Pt(13) if line == 'PROJE RAPORU' else Pt(12)
    rr.bold = (line == 'PROJE RAPORU')
    _set_font(rr)

doc.add_page_break()

# ======================== İÇİNDEKİLER ========================
h1('İçindekiler')
icindekiler = [
    '1. Özet', '2. Proje Tanımı', '3. Gereksinim Analizi', '4. Kullanım Senaryoları',
    '5. Sistem Mimarisi', '6. Kullanılan Teknolojiler', '7. Veri Modeli',
    '8. Gerçekleştirilen Modüller', '9. API Açıklamaları', '10. Kurulum Adımları',
    '11. Ekip İçi Görev Dağılımı', '12. Test Süreci ve Test Senaryoları',
    '13. Karşılaşılan Kısıtlar', '14. Sonuç ve Gelecek Çalışmalar', 'Ek: Şekil Listesi',
]
for i in icindekiler:
    para(i, justify=False)
doc.add_page_break()

# ============================ 1. ÖZET ============================
h1('1. Özet')
para('SafeDrive, lojistik şirketlerinin araç filolarını ve sürücülerini gerçek zamanlı '
     'olarak izlemesini sağlayan uçtan uca bir telematik platformudur. Sistem; sürücülerin '
     'akıllı telefonlarından gelen ivmeölçer, jiroskop ve GPS verilerini toplar, bu verileri '
     'analiz ederek riskli sürüş davranışlarını (ani fren, sert dönüş, ani hızlanma, sarsıntı) '
     'otomatik olarak tespit eder ve bağlı şirket paneline anlık olarak iletir. Platform üç '
     'farklı kullanıcı rolüne (yönetici, lojistik şirketi, sürücü) hizmet verir.')
para('Proje; Node.js/Express tabanlı bir REST API ve Socket.io gerçek zamanlı iletişim '
     'katmanı, React (Vite) tabanlı bir web yönetim paneli ve tarayıcı tabanlı bir mobil '
     'sürücü uygulamasından oluşur. Veriler MongoDB üzerinde tutulur. Sistem bulut ortamına '
     '(MongoDB Atlas, Render, Vercel) deploy edilerek internet üzerinden erişilebilir hale '
     'getirilmiştir.')

# ============================ 2. PROJE TANIMI ============================
h1('2. Proje Tanımı')
para('Lojistik sektöründe sürücü güvenliği ve filo verimliliği kritik öneme sahiptir. '
     'Geleneksel takip sistemleri yalnızca konum bilgisi sunarken, SafeDrive sürücü '
     'davranışını da ölçerek riskli durumları proaktif olarak tespit eder.')
h2('2.1. Amaç')
bullet('Sürücülerin telefon sensörlerinden gelen verilerle riskli davranışları otomatik tespit etmek.')
bullet('Lojistik şirketlerine filolarını canlı harita üzerinden izleme imkânı sunmak.')
bullet('Anomali (alarm) durumlarını saniyeler içinde ilgili panele iletmek.')
bullet('Yöneticilere kullanıcı, cihaz ve şirket yönetimi için merkezi bir panel sağlamak.')
h2('2.2. Kapsam')
para('Sistem üç ana arayüzden oluşur: (1) yöneticinin tüm sistemi yönettiği web paneli, '
     '(2) lojistik şirketinin kendi filosunu izlediği harita tabanlı panel, (3) sürücünün '
     'sensör verisi gönderdiği mobil web uygulaması. Ayrıca yeni şirketlerin başvuru '
     'gönderebildiği bir tanıtım (landing) sayfası bulunur.')
sekil('Açılış (landing) sayfası — tanıtım bölümü ve en alttaki başvuru formu')

# ============================ 3. GEREKSİNİM ANALİZİ ============================
h1('3. Gereksinim Analizi')
h2('3.1. Fonksiyonel Gereksinimler')
tablo(['No', 'Gereksinim'], [
    ['FR-01', 'Kullanıcılar rol bazlı (yönetici/şirket/sürücü) kayıt olabilmeli ve giriş yapabilmeli.'],
    ['FR-02', 'Sürücü mobil uygulamadan ivmeölçer, jiroskop ve GPS verisi gönderebilmeli.'],
    ['FR-03', 'Sistem gelen veriyi analiz edip anomali (ani fren, sert dönüş, ani hızlanma, sarsıntı) tespit etmeli.'],
    ['FR-04', 'Tespit edilen anomaliler alarm olarak kaydedilmeli ve anlık bildirilmeli.'],
    ['FR-05', 'Şirket yalnızca kendi araçlarının verisini ve konumunu canlı harita üzerinde görebilmeli.'],
    ['FR-06', 'Yönetici tüm kullanıcı ve cihazları görüntüleyip düzenleyip silebilmeli.'],
    ['FR-07', 'Yeni şirketler tanıtım sayfasından başvuru gönderebilmeli; başvurular yöneticiye düşmeli.'],
    ['FR-08', 'Veriler ve konum güncellemeleri panele gerçek zamanlı (Socket.io) yansımalı.'],
])
h2('3.2. Fonksiyonel Olmayan Gereksinimler')
bullet('Güvenlik: Parolalar bcrypt ile hash’lenir; yetkilendirme JWT ile yapılır; veri erişimi role göre kısıtlanır.')
bullet('Erişilebilirlik: Sistem internet üzerinden (bulut) her cihazdan erişilebilir olmalıdır.')
bullet('Performans: Sensör verisi 2 saniyelik aralıklarla işlenmeli ve anlık iletilmelidir.')
bullet('Kullanılabilirlik: Arayüz tamamen Türkçe, sade ve mobil uyumlu olmalıdır.')
bullet('Taşınabilirlik: Mobil uygulama ek kurulum gerektirmeden tarayıcıdan çalışmalıdır.')

# ============================ 4. KULLANIM SENARYOLARI ============================
h1('4. Kullanım Senaryoları')
h2('4.1. Aktörler')
bullet('Yönetici (admin): Tüm sistemi yöneten kişi.')
bullet('Lojistik Şirketi (company): Kendi filosunu izleyen kurumsal kullanıcı.')
bullet('Sürücü (driver): Mobil uygulamadan veri üreten saha kullanıcısı.')
h2('4.2. Senaryo — Sürücü Veri Toplama')
numbered('Sürücü mobil web uygulamasını telefon tarayıcısında açar.')
numbered('Şirketini seçerek kayıt olur veya giriş yapar; sistem ona otomatik bir cihaz tanımlar.')
numbered('"Başlat" butonuna basar; hareket ve konum izinlerini onaylar.')
numbered('Uygulama her 2 saniyede sensör + konum verisini sunucuya gönderir.')
h2('4.3. Senaryo — Şirketin Filoyu İzlemesi')
numbered('Şirket kullanıcısı kendi paneline giriş yapar.')
numbered('Sol tarafta araç listesini (sürücü, son görülme, çevrimiçi/çevrimdışı) görür.')
numbered('Sağdaki harita üzerinde araçların canlı konumunu izler.')
numbered('Bir araçta anomali oluşunca harita ve alarm listesi anında güncellenir, uyarı çıkar.')
h2('4.4. Senaryo — Yönetici Yönetimi')
numbered('Yönetici panele giriş yapar.')
numbered('Kullanıcıları, cihazları ve şirket başvurularını görüntüler.')
numbered('Gerekirse kullanıcı/cihaz düzenler, siler veya şirket atamasını değiştirir.')
sekil('Şirket paneli — solda araç listesi, sağda Leaflet haritası üzerinde canlı araç konumu')

# ============================ 5. SİSTEM MİMARİSİ ============================
h1('5. Sistem Mimarisi')
para('SafeDrive üç katmanlı (istemci – sunucu – veritabanı) bir mimariye sahiptir. İstemci '
     'katmanı web paneli ve mobil uygulamadan; sunucu katmanı REST API ve Socket.io '
     'gerçek zamanlı iletişim katmanından; veri katmanı ise MongoDB veritabanından oluşur.')
h2('5.1. Bileşenler')
tablo(['Katman', 'Bileşen', 'Görev'], [
    ['İstemci', 'React Web Paneli', 'Yönetici ve şirket arayüzleri'],
    ['İstemci', 'Mobil Web Uygulaması', 'Sürücü sensör verisi toplama'],
    ['Sunucu', 'Express REST API', 'Kimlik doğrulama, veri kaydı, iş kuralları'],
    ['Sunucu', 'Socket.io', 'Gerçek zamanlı veri ve alarm yayını'],
    ['Sunucu', 'Anomali Tespit Servisi', 'Sürüş davranışı analizi'],
    ['Veri', 'MongoDB (Atlas)', 'Kalıcı veri saklama'],
])
h2('5.2. Veri Akışı')
numbered('Mobil uygulama sensör verisini JWT ile korunan POST /api/sensor-data adresine gönderir.')
numbered('Sunucu veriyi kaydeder ve Anomali Tespit Servisi’ni çağırır.')
numbered('Anomali varsa bir Alarm kaydı oluşturulur.')
numbered('Socket.io ile "newData" (ve gerekiyorsa "newAlarm") olayı tüm panellere yayınlanır.')
numbered('İlgili şirket/yönetici paneli arayüzü anlık olarak güncellenir.')
sekil('Sistem mimarisi şeması — İstemci (web + mobil) → Express API + Socket.io → MongoDB akış diyagramı')
h2('5.3. Dağıtım (Deployment) Mimarisi')
tablo(['Bileşen', 'Platform', 'Açıklama'], [
    ['Veritabanı', 'MongoDB Atlas', 'Bulut veritabanı (ücretsiz M0 cluster)'],
    ['Backend', 'Render', 'Node.js web servisi, HTTPS'],
    ['Frontend', 'Vercel', 'Statik React paneli, HTTPS'],
    ['Mobil', 'Render (backend ile)', 'Backend tarafından sunulan statik HTML'],
])
sekil('Dağıtım şeması — Atlas + Render + Vercel bağlantısını gösteren bulut mimarisi diyagramı')

# ============================ 6. KULLANILAN TEKNOLOJİLER ============================
h1('6. Kullanılan Teknolojiler')
h2('6.1. Backend')
for x in ['Node.js & Express.js — REST API sunucusu', 'MongoDB & Mongoose — veritabanı ve ODM',
          'JSON Web Token (JWT) — kimlik doğrulama', 'bcrypt — parola hash’leme',
          'Socket.io — gerçek zamanlı iletişim', 'express-validator — girdi doğrulama',
          'cors, dotenv — yapılandırma ve erişim']:
    bullet(x)
h2('6.2. Frontend (Web Paneli)')
for x in ['React (Vite) — kullanıcı arayüzü', 'React Router DOM — sayfa yönlendirme',
          'Axios — HTTP istekleri (token interceptor’lı)', 'Socket.io-client — canlı veri alımı',
          'Chart.js / react-chartjs-2 — ivme grafiği',
          'Leaflet / react-leaflet — canlı harita ve araç işaretçileri',
          'react-hot-toast — bildirimler', 'jwt-decode — token çözümleme']:
    bullet(x)
h2('6.3. Mobil Uygulama')
for x in ['Saf HTML/CSS/JavaScript (kütüphanesiz)', 'DeviceMotion API — ivmeölçer ve jiroskop',
          'Geolocation API — GPS konumu']:
    bullet(x)
h2('6.4. Dağıtım ve Araçlar')
for x in ['MongoDB Atlas, Render, Vercel — bulut servisleri', 'Git & GitHub — sürüm kontrolü',
          'Docker — yerel MongoDB geliştirme ortamı']:
    bullet(x)

# ============================ 7. VERİ MODELİ ============================
h1('7. Veri Modeli')
para('Sistem altı ana koleksiyondan (veri modelinden) oluşur. Aşağıda her birinin alanları verilmiştir.')
h2('7.1. User (Kullanıcı)')
tablo(['Alan', 'Tip', 'Açıklama'], [
    ['username', 'String', 'Benzersiz kullanıcı adı'], ['email', 'String', 'Benzersiz e-posta'],
    ['password', 'String', 'bcrypt ile hash’lenmiş parola'], ['role', 'String', "enum: 'admin' | 'driver'"],
    ['companyId', 'ObjectId', 'Bağlı olduğu şirket (sürücüler için)'], ['createdAt', 'Date', 'Kayıt tarihi'],
])
h2('7.2. Company (Şirket)')
tablo(['Alan', 'Tip', 'Açıklama'], [
    ['name', 'String', 'Benzersiz şirket adı'], ['email', 'String', 'Benzersiz e-posta'],
    ['password', 'String', 'bcrypt ile hash’lenmiş parola'], ['createdAt', 'Date', 'Kayıt tarihi'],
])
h2('7.3. Device (Cihaz / Araç)')
tablo(['Alan', 'Tip', 'Açıklama'], [
    ['deviceId', 'String', 'Benzersiz cihaz kimliği'], ['owner', 'ObjectId', 'Sahip sürücü (User)'],
    ['companyId', 'ObjectId', 'Bağlı şirket (Company)'], ['model', 'String', 'Cihaz modeli'],
    ['platform', 'String', "enum: 'android' | 'ios' | 'web'"], ['registeredAt', 'Date', 'Kayıt tarihi'],
    ['lastSeen', 'Date', 'Son veri zamanı'],
])
h2('7.4. SensorData (Sensör Verisi)')
tablo(['Alan', 'Tip', 'Açıklama'], [
    ['deviceId', 'ObjectId', 'Veriyi gönderen cihaz'], ['timestamp', 'Date', 'Zaman damgası (indeksli)'],
    ['accelerometer', '{x,y,z}', 'İvmeölçer değerleri'], ['gyroscope', '{alpha,beta,gamma}', 'Jiroskop değerleri'],
    ['location', '{latitude,longitude,speed}', 'GPS konumu ve hız'],
])
h2('7.5. Alarm')
tablo(['Alan', 'Tip', 'Açıklama'], [
    ['deviceId', 'ObjectId', 'İlgili cihaz'],
    ['type', 'String', 'HARD_BRAKE | SHARP_TURN | RAPID_ACCELERATION | VIBRATION'],
    ['severity', 'String', 'low | medium | high | critical'], ['value', 'Number', 'Tetikleyen ölçüm değeri'],
    ['timestamp', 'Date', 'Oluşma zamanı'], ['resolved', 'Boolean', 'Çözüldü mü'],
])
h2('7.6. Application (Başvuru)')
tablo(['Alan', 'Tip', 'Açıklama'], [
    ['companyName', 'String', 'Başvuran şirket adı'], ['contactName', 'String', 'Yetkili kişi'],
    ['email', 'String', 'İletişim e-postası'], ['phone', 'String', 'Telefon'],
    ['message', 'String', 'Mesaj'], ['reviewed', 'Boolean', 'İncelendi mi'],
    ['createdAt', 'Date', 'Başvuru tarihi'],
])
sekil('Veri modeli ilişki şeması (ER diyagramı) — User, Company, Device, SensorData, Alarm ilişkileri')

# ============================ 8. MODÜLLER ============================
h1('8. Gerçekleştirilen Modüller')
h2('8.1. Kimlik Doğrulama ve Yetkilendirme')
para('JWT tabanlı kimlik doğrulama; bcrypt ile parola hash’leme; rol bazlı erişim (admin/driver/company). '
     'Her istekte Authorization: Bearer <token> başlığı doğrulanır.')
h2('8.2. Sensör Veri Toplama')
para('Mobil uygulama DeviceMotion ve Geolocation API’leri ile veri toplayıp 2 saniyede bir API’ye gönderir. '
     'Cihaz kimliği otomatik atanır (sürücü manuel işlem yapmaz).')
h2('8.3. Anomali Tespiti')
para('Kaydedilen her veride aşağıdaki kurallar çalıştırılır:')
tablo(['Anomali', 'Koşul', 'Şiddet'], [
    ['HARD_BRAKE (Ani Fren)', 'accelerometer.x < -8', 'critical'],
    ['RAPID_ACCELERATION (Ani Hızlanma)', 'accelerometer.x > 10', 'high'],
    ['SHARP_TURN (Sert Dönüş)', 'gyroscope.gamma > 150 veya < -150', 'high'],
    ['VIBRATION (Sarsıntı)', 'son 5 verinin accelerometer.x standart sapması > 4', 'medium'],
])
h2('8.4. Alarm Yönetimi')
para('Anomaliler Alarm olarak kaydedilir; yönetici alarmları "çözüldü" işaretleyebilir; '
     'şirket ve sürücü yalnızca kendi alarmlarını görür.')
h2('8.5. Gerçek Zamanlı İletişim')
para('Socket.io ile "newData" ve "newAlarm" olayları tüm bağlı istemcilere yayınlanır; '
     'paneller sayfa yenilemeden güncellenir.')
h2('8.6. Filo Takip ve Harita')
para('Şirket paneli Leaflet/OpenStreetMap üzerinde araçları renkli işaretçilerle (çevrimiçi=yeşil, '
     'çevrimdışı=gri) gösterir; işaretçiye tıklanınca sürücü, hız ve son alarm bilgisi açılır.')
h2('8.7. Yönetici Yönetim Modülü')
para('Kullanıcı ve cihazlar üzerinde tam CRUD: düzenleme (rol/şirket/sahip değiştirme) ve silme '
     '(ilişkili verilerle birlikte) işlemleri.')
h2('8.8. Başvuru (Landing) Modülü')
para('Tanıtım sayfasındaki form ile şirketler başvuru gönderir; başvurular yönetici panelinde listelenir.')
sekil('Yönetici paneli — Dashboard: canlı ivme grafiği (Chart.js) ve özet kartlar')
sekil('Yönetici paneli — Kullanıcılar veya Cihazlar sayfası (düzenle/sil işlemleriyle)')
sekil('Mobil uygulama — sürücü sensör ekranı (ivmeölçer, jiroskop, GPS değerleri)')

# ============================ 9. API ============================
h1('9. API Açıklamaları')
para('Tüm yanıtlar JSON formatındadır. Korumalı uç noktalar Authorization: Bearer <token> başlığı gerektirir.')
h2('9.1. Kimlik Doğrulama')
tablo(['Metot', 'Uç Nokta', 'Açıklama'], [
    ['POST', '/auth/register', 'Sürücü kaydı (username, email, password, companyId)'],
    ['POST', '/auth/login', 'Yönetici/sürücü girişi (email, password)'],
    ['POST', '/auth/register-company', 'Şirket kaydı (name, email, password)'],
    ['POST', '/auth/login-company', 'Şirket girişi (email, password)'],
])
h2('9.2. Şirketler ve Sensör')
tablo(['Metot', 'Uç Nokta', 'Yetki', 'Açıklama'], [
    ['GET', '/api/companies', 'Açık', 'Şirket listesi (kayıt formu için)'],
    ['POST', '/api/sensor-data', 'JWT', 'Sensör verisi kaydı + anomali analizi'],
    ['GET', '/api/sensor-data', 'JWT', 'Sensör verisi sorgulama (rol bazlı)'],
])
h2('9.3. Alarmlar ve Cihazlar')
tablo(['Metot', 'Uç Nokta', 'Yetki', 'Açıklama'], [
    ['GET', '/api/alarms', 'JWT', 'Alarm listesi (rol bazlı filtre)'],
    ['PATCH', '/api/alarms/:id', 'Admin', 'Alarmı çözüldü işaretle'],
    ['GET', '/api/devices', 'JWT', 'Cihaz listesi (rol bazlı)'],
    ['POST', '/api/devices', 'JWT', 'Yeni cihaz kaydı'],
    ['PATCH', '/api/devices/:id', 'Admin', 'Cihaz güncelleme'],
    ['DELETE', '/api/devices/:id', 'Admin', 'Cihaz silme (ilişkili verilerle)'],
])
h2('9.4. Kullanıcılar ve Başvurular')
tablo(['Metot', 'Uç Nokta', 'Yetki', 'Açıklama'], [
    ['GET', '/api/users', 'Admin', 'Kullanıcı listesi'],
    ['PATCH', '/api/users/:id', 'Admin', 'Kullanıcı güncelleme (rol/şirket)'],
    ['DELETE', '/api/users/:id', 'Admin', 'Kullanıcı silme'],
    ['POST', '/api/applications', 'Açık', 'Başvuru gönderme'],
    ['GET', '/api/applications', 'Admin', 'Başvuru listesi'],
    ['PATCH', '/api/applications/:id', 'Admin', 'Başvuru incelendi işaretle'],
])
h2('9.5. Örnek İstek/Yanıt — POST /api/sensor-data')
code('İSTEK gövdesi (Body):\n'
     '{\n'
     '  "deviceId": "<cihaz_objectid>",\n'
     '  "timestamp": "2026-06-07T10:00:00.000Z",\n'
     '  "accelerometer": { "x": -9, "y": 0.1, "z": 9.8 },\n'
     '  "gyroscope": { "alpha": 0, "beta": 0, "gamma": 0 },\n'
     '  "location": { "latitude": 41.01, "longitude": 28.97, "speed": 20 }\n'
     '}')
code('YANIT (201):\n'
     '{\n'
     '  "sensorData": { ... },\n'
     '  "alarm": { "type": "HARD_BRAKE", "severity": "critical", "value": -9 }\n'
     '}')
sekil('API testi — bir uç noktanın Postman veya tarayıcı çıktısı (JSON yanıt)')

# ============================ 10. KURULUM ============================
h1('10. Kurulum Adımları')
h2('10.1. Yerel Geliştirme Ortamı')
para('Gereksinimler: Node.js, Docker (veya yerel MongoDB), Git.')
numbered('Depoyu klonlayın: git clone <repo-adresi>')
numbered('MongoDB’yi başlatın: docker run -d --name safedrive-mongo -p 27017:27017 mongo:7')
numbered('Backend bağımlılıkları: cd safedrive-backend && npm install')
numbered('.env dosyasını oluşturun (MONGO_URI, JWT_SECRET, PORT).')
numbered('Backend’i çalıştırın: npm run dev')
numbered('Frontend bağımlılıkları: cd safedrive-frontend && npm install')
numbered('Frontend’i çalıştırın: npm run dev (http://localhost:5173)')
h2('10.2. Backend .env Örneği')
code('MONGO_URI=mongodb://localhost:27017/safedrive\n'
     'JWT_SECRET=gizli_anahtar\n'
     'PORT=5000')
h2('10.3. Bulut Dağıtımı (Production)')
numbered('MongoDB Atlas’ta ücretsiz cluster oluşturulur; kullanıcı ve ağ erişimi (0.0.0.0/0) ayarlanır.')
numbered('Backend Render’a deploy edilir (Root: safedrive-backend, Build: npm install, Start: npm start).')
numbered('Render ortam değişkenleri: MONGO_URI (Atlas), JWT_SECRET.')
numbered('Frontend Vercel’e deploy edilir (Root: safedrive-frontend, env: VITE_API_URL = Render backend adresi).')
numbered('Mobil uygulama backend ile birlikte sunulur: <backend-adresi>/mobile.html')
sekil('Dağıtım kanıtı — Render başarılı deploy logu ("Your service is live") veya Vercel deploy ekranı')

# ============================ 11. GÖREV DAĞILIMI ============================
h1('11. Ekip İçi Görev Dağılımı')
para('[DÜZENLE: Aşağıdaki tabloyu kendi ekip üyelerinize göre güncelleyin. Tek kişilik projede '
     'tek satır bırakabilirsiniz.]')
tablo(['Ekip Üyesi', 'Sorumlu Olduğu Bölümler'], [
    ['[Ad Soyad 1]', 'Backend API, veri modeli, kimlik doğrulama'],
    ['[Ad Soyad 2]', 'Anomali tespiti, Socket.io gerçek zamanlı katman'],
    ['[Ad Soyad 3]', 'Web paneli (React), grafik ve harita arayüzleri'],
    ['[Ad Soyad 4]', 'Mobil uygulama, test ve bulut dağıtımı'],
])

# ============================ 12. TEST ============================
h1('12. Test Süreci ve Test Senaryoları')
para('Sistem, geliştirme sürecinde uç nokta (API) testleri ve uçtan uca senaryo testleri ile '
     'doğrulanmıştır. Aşağıda yürütülen başlıca test senaryoları ve sonuçları verilmiştir.')
tablo(['No', 'Senaryo', 'Beklenen Sonuç', 'Durum'], [
    ['TC-01', 'Geçerli bilgilerle yönetici girişi', 'JWT token döner, panele erişilir', 'Başarılı'],
    ['TC-02', 'Hatalı parola ile giriş', '401 Yetkisiz hatası', 'Başarılı'],
    ['TC-03', 'Token’sız korumalı uç noktaya erişim', '401 hatası', 'Başarılı'],
    ['TC-04', 'Ani fren verisi gönderimi (x=-9)', 'HARD_BRAKE / critical alarmı oluşur', 'Başarılı'],
    ['TC-05', 'Ani hızlanma verisi (x=12)', 'RAPID_ACCELERATION / high alarmı', 'Başarılı'],
    ['TC-06', 'Sert dönüş verisi (gamma=160)', 'SHARP_TURN / high alarmı', 'Başarılı'],
    ['TC-07', 'Sürücü kaydında şirket seçilmemesi', '400 "Şirket seçimi zorunludur"', 'Başarılı'],
    ['TC-08', 'Şirketin başka şirketin verisine erişimi', 'Yalnızca kendi araçlarını görür', 'Başarılı'],
    ['TC-09', 'Yöneticinin cihaz silmesi', 'Cihaz ve ilişkili veriler silinir', 'Başarılı'],
    ['TC-10', 'Yeni veride panelin canlı güncellenmesi', 'Socket.io ile anlık güncelleme', 'Başarılı'],
    ['TC-11', 'Mobil uygulamadan canlı konum gönderimi', 'Harita üzerinde araç işaretçisi belirir', 'Başarılı'],
    ['TC-12', 'Landing formundan başvuru gönderimi', 'Başvuru kaydedilir, yöneticiye düşer', 'Başarılı'],
])
sekil('Test kanıtı — bir alarmın panelde canlı belirmesi (toast bildirimi) veya API test çıktısı')

# ============================ 13. KISITLAR ============================
h1('13. Karşılaşılan Kısıtlar')
bullet('Tarayıcı güvenlik kısıtı: DeviceMotion ve Geolocation API’leri yalnızca HTTPS (veya localhost) '
       'üzerinde çalışır. Bu nedenle mobil test için HTTPS gereklidir; bulut dağıtımı (Render HTTPS) ile çözülmüştür.')
bullet('GPS izni: Sürücünün konum iznini reddetmesi durumunda konum verisi boş gelir; uygulamaya görünür '
       'hata mesajı eklenerek bu durum yönetilmiştir.')
bullet('Render ücretsiz plan: Servis 15 dakika işlemsizlikte uykuya geçer; ilk istek ~50 saniye gecikir.')
bullet('GPS doğruluğu: Dizüstü/masaüstü tarayıcılarda GPS donanımı olmadığından konum yaklaşık alınır; '
       'doğru test için gerçek mobil cihaz gerekir.')
bullet('Tek bölge veritabanı: Ücretsiz Atlas cluster tek bölgede çalışır; yüksek ölçekli üretim için '
       'yeterli değildir.')

# ============================ 14. SONUÇ ============================
h1('14. Sonuç ve Gelecek Çalışmalar')
para('SafeDrive projesi ile lojistik filolar için gerçek zamanlı sürücü davranış analizi ve filo takip '
     'sistemi başarıyla gerçekleştirilmiştir. Sistem; sensör verisi toplama, anomali tespiti, gerçek '
     'zamanlı bildirim, rol bazlı yetkilendirme ve harita tabanlı izleme yeteneklerini uçtan uca '
     'sağlamakta ve bulut ortamında çalışmaktadır.')
para('Gelecek çalışmalar olarak şunlar planlanabilir:')
bullet('Sürücülerin geçmiş rotasının harita üzerinde çizgi olarak gösterilmesi.')
bullet('Sürücü risk skorlaması ve haftalık/aylık raporlama.')
bullet('Mobil uygulamanın native (Android/iOS) sürümü ile arka plan veri toplama.')
bullet('Makine öğrenmesi ile daha gelişmiş anomali tespiti.')
bullet('E-posta/SMS ile kritik alarm bildirimi.')

# ============================ EK ============================
doc.add_page_break()
h1('Ek: Şekil Listesi')
para('Aşağıdaki şekiller raporun ilgili bölümlerine eklenmelidir.')
for n, acik in sekil_listesi:
    p = doc.add_paragraph(style='List Number')
    p.add_run(f'Şekil {n}: ').bold = True
    p.add_run(acik)

doc.save('/home/nihat/Desktop/NodeJS/SafeDrive_Rapor.docx')
print('Rapor olusturuldu. Sekil sayisi:', sekil_sayac[0])
